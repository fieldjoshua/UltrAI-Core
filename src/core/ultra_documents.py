import hashlib
import io
import logging
import os
import pickle
import re
import time
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, Union

import faiss
import fitz  # PyMuPDF
import numpy as np
from langchain.text_splitter import RecursiveCharacterTextSplitter
from sentence_transformers import SentenceTransformer
from ultra_base import UltraBase


class DocumentCache:
    """Cache for document processing to avoid redundant operations."""

    def __init__(self, cache_dir: str = ".cache"):
        """Initialize the document cache."""
        self.cache_dir = Path(cache_dir)
        self.cache_dir.mkdir(exist_ok=True)
        self.logger = logging.getLogger(__name__)

    def _get_cache_path(self, file_path: str, operation: str) -> Path:
        """Generate a unique cache path for a file and operation."""
        file_hash = hashlib.md5(file_path.encode()).hexdigest()
        file_timestamp = os.path.getmtime(file_path)
        cache_key = f"{file_hash}_{int(file_timestamp)}_{operation}"
        return self.cache_dir / f"{cache_key}.pkl"

    def get(self, file_path: str, operation: str) -> Optional[Any]:
        """Get cached result for a file and operation if available."""
        if not os.path.exists(file_path):
            return None

        cache_path = self._get_cache_path(file_path, operation)
        if cache_path.exists():
            try:
                with open(cache_path, "rb") as f:
                    return pickle.load(f)
            except Exception as e:
                self.logger.warning(f"Failed to load cache for {file_path}: {e}")
                return None
        return None

    def set(self, file_path: str, operation: str, data: Any) -> bool:
        """Cache result for a file and operation."""
        if not os.path.exists(file_path):
            return False

        cache_path = self._get_cache_path(file_path, operation)
        try:
            with open(cache_path, "wb") as f:
                pickle.dump(data, f)
            return True
        except Exception as e:
            self.logger.warning(f"Failed to save cache for {file_path}: {e}")
            return False


class UltraDocuments(UltraBase):
    def __init__(
        self,
        api_keys: Dict[str, str],
        prompt_templates: Optional[Any] = None,
        rate_limits: Optional[Any] = None,
        output_format: str = "plain",
        enabled_features: Optional[List[str]] = None,
        cache_enabled: bool = True,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        embedding_model: str = "all-MiniLM-L6-v2",
    ):
        super().__init__(
            api_keys=api_keys,
            prompt_templates=prompt_templates,
            rate_limits=rate_limits,
            output_format=output_format,
            enabled_features=enabled_features,
        )

        # File processing handlers
        self.supported_formats = {
            ".pdf": self._process_pdf,
            ".txt": self._process_text,
            ".md": self._process_text,
            ".docx": self._process_docx,
            ".doc": self._process_doc,
        }

        # Document processing settings
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.cache_enabled = cache_enabled

        # Initialize cache
        if self.cache_enabled:
            self.cache = DocumentCache()
            self.logger.info("Document cache initialized")

        # Initialize text splitter for chunking
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            length_function=len,
            is_separator_regex=False,
        )

        # Initialize embedding model
        try:
            self.embedding_model = SentenceTransformer(embedding_model)
            self.embedding_dimension = (
                self.embedding_model.get_sentence_embedding_dimension()
            )
            self.index = None  # Will be initialized when needed
            self.embeddings = []
            self.chunks = []
            self.doc_mapping = {}  # Maps chunk index to document source
            self.logger.info(f"Document embedding model initialized: {embedding_model}")
        except Exception as e:
            self.logger.error(f"Failed to initialize embedding model: {e}")
            self.embedding_model = None

    def _process_pdf(self, file_path: Union[str, Path]) -> str:
        """Extract text from a PDF file using PyMuPDF (faster and more reliable than PyPDF2)."""
        try:
            # Check cache first
            if self.cache_enabled:
                cached_text = self.cache.get(str(file_path), "pdf_text")
                if cached_text is not None:
                    self.logger.info(f"Retrieved PDF text from cache: {file_path}")
                    return cached_text

            # Process PDF if not in cache
            start_time = time.time()
            text = ""
            with fitz.open(file_path) as doc:
                for page_num, page in enumerate(doc):
                    text += page.get_text() + "\n\n"

            # Cache the result
            if self.cache_enabled:
                self.cache.set(str(file_path), "pdf_text", text)

            processing_time = time.time() - start_time
            self.logger.info(f"Processed PDF in {processing_time:.2f}s: {file_path}")
            return text

        except Exception as e:
            self.logger.error(f"Error processing PDF file {file_path}: {str(e)}")
            raise

    def _process_text(self, file_path: Union[str, Path]) -> str:
        """Extract text from a text file."""
        try:
            with open(file_path, "r", encoding="utf-8") as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encodings if UTF-8 fails
            encodings = ["latin-1", "cp1252", "iso-8859-1"]
            for encoding in encodings:
                try:
                    with open(file_path, "r", encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
            self.logger.error(
                f"Could not decode text file {file_path} with any encoding"
            )
            raise
        except Exception as e:
            self.logger.error(f"Error processing text file {file_path}: {str(e)}")
            raise

    def _process_docx(self, file_path: Union[str, Path]) -> str:
        """Extract text from a DOCX file."""
        try:
            # Check cache first
            if self.cache_enabled:
                cached_text = self.cache.get(str(file_path), "docx_text")
                if cached_text is not None:
                    self.logger.info(f"Retrieved DOCX text from cache: {file_path}")
                    return cached_text

            from docx import Document

            doc = Document(file_path)

            # Extract paragraphs with better formatting preservation
            text = []
            for para in doc.paragraphs:
                if para.text.strip():  # Skip empty paragraphs
                    text.append(para.text)

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text.append(" | ".join(row_text))

            result = "\n\n".join(text)

            # Cache the result
            if self.cache_enabled:
                self.cache.set(str(file_path), "docx_text", result)

            return result

        except Exception as e:
            self.logger.error(f"Error processing DOCX file {file_path}: {str(e)}")
            raise

    def _process_doc(self, file_path: Union[str, Path]) -> str:
        """Extract text from a DOC file."""
        try:
            # Check cache first
            if self.cache_enabled:
                cached_text = self.cache.get(str(file_path), "doc_text")
                if cached_text is not None:
                    self.logger.info(f"Retrieved DOC text from cache: {file_path}")
                    return cached_text

            import textract

            text = textract.process(file_path).decode("utf-8")

            # Cache the result
            if self.cache_enabled:
                self.cache.set(str(file_path), "doc_text", text)

            return text

        except Exception as e:
            self.logger.error(f"Error processing DOC file {file_path}: {str(e)}")
            raise

    def process_document(self, file_path: Union[str, Path]) -> str:
        """Process a document and extract its text content."""
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        file_extension = file_path.suffix.lower()
        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}")

        return self.supported_formats[file_extension](file_path)

    def process_documents(self, file_paths: List[Union[str, Path]]) -> Dict[str, str]:
        """Process multiple documents and return their text content."""
        results = {}
        for file_path in file_paths:
            try:
                results[str(file_path)] = self.process_document(file_path)
            except Exception as e:
                self.logger.error(f"Error processing {file_path}: {str(e)}")
                results[str(file_path)] = f"Error: {str(e)}"
        return results

    def chunk_text(self, text: str) -> List[str]:
        """Split text into manageable chunks."""
        return self.text_splitter.split_text(text)

    def chunk_documents(self, documents: Dict[str, str]) -> List[Tuple[str, str]]:
        """Split documents into chunks and maintain mapping to source."""
        chunked_docs = []
        for doc_path, text in documents.items():
            chunks = self.chunk_text(text)
            for chunk in chunks:
                chunked_docs.append((doc_path, chunk))
        return chunked_docs

    def embed_text(self, text: str) -> np.ndarray:
        """Generate embedding for a text string."""
        if self.embedding_model is None:
            raise ValueError("Embedding model not initialized")
        return self.embedding_model.encode(text)

    def embed_chunks(self, chunks: List[Tuple[str, str]]) -> None:
        """Embed all chunks and prepare for retrieval."""
        if self.embedding_model is None:
            self.logger.error("Embedding model not initialized, skipping embedding")
            return

        self.chunks = [chunk[1] for chunk in chunks]
        self.doc_mapping = {i: chunks[i][0] for i in range(len(chunks))}

        # Generate embeddings for all chunks
        self.embeddings = self.embedding_model.encode(self.chunks)

        # Create FAISS index for fast retrieval
        self.index = faiss.IndexFlatL2(self.embedding_dimension)
        faiss.normalize_L2(self.embeddings)
        self.index.add(self.embeddings)

        self.logger.info(
            f"Created index with {len(self.chunks)} chunks from {len(set(self.doc_mapping.values()))} documents"
        )

    def retrieve_relevant_chunks(
        self, query: str, top_k: int = 5
    ) -> List[Tuple[str, str, float]]:
        """Retrieve the most relevant chunks for a query."""
        if self.index is None or not self.chunks:
            raise ValueError("No documents have been indexed yet")

        # Embed the query
        query_embedding = self.embedding_model.encode([query])
        faiss.normalize_L2(query_embedding)

        # Search for similar chunks
        distances, indices = self.index.search(query_embedding, top_k)

        # Return relevant chunks with metadata
        results = []
        for i, idx in enumerate(indices[0]):
            if idx < len(self.chunks):  # Safety check
                doc_path = self.doc_mapping[idx]
                doc_name = os.path.basename(doc_path)
                results.append((doc_name, self.chunks[idx], float(distances[0][i])))

        return results

    def process_documents_with_embedding(
        self, file_paths: List[Union[str, Path]]
    ) -> List[Tuple[str, str, float]]:
        """Process documents, chunk them, and create embeddings for retrieval."""
        # Process documents
        documents = self.process_documents(file_paths)

        # Chunk documents
        chunked_docs = self.chunk_documents(documents)

        # Embed chunks
        self.embed_chunks(chunked_docs)

        return [
            (os.path.basename(doc_path), chunk, 0.0) for doc_path, chunk in chunked_docs
        ]

    def process_query_with_documents(
        self, query: str, file_paths: List[Union[str, Path]], max_chunks: int = 5
    ) -> str:
        """
        Process a query with documents to create a context-enhanced prompt.

        This is a simplified RAG (Retrieval Augmented Generation) approach.
        """
        if not file_paths:
            return query

        # Process and index documents if needed
        if self.index is None or not self.chunks:
            self.process_documents_with_embedding(file_paths)

        # Get relevant chunks (if embedding model is available)
        if self.embedding_model is not None:
            try:
                relevant_chunks = self.retrieve_relevant_chunks(query, top_k=max_chunks)
                context_parts = []

                for doc_name, chunk, score in relevant_chunks:
                    context_parts.append(
                        f"[From: {doc_name}, Relevance: {100 * (1 - score):.1f}%]\n{chunk}"
                    )

                context = "\n\n".join(context_parts)
                enhanced_prompt = (
                    f"{query}\n\n"
                    f"--- RELEVANT DOCUMENT EXCERPTS ---\n{context}\n"
                    f"--- END OF DOCUMENT EXCERPTS ---\n\n"
                    f"Please provide a comprehensive response to the original query, "
                    f"using these document excerpts as context."
                )
                return enhanced_prompt

            except Exception as e:
                self.logger.error(
                    f"Error in semantic retrieval: {e}. Falling back to basic document processing."
                )

        # Fallback: traditional approach (concatenate all document contents)
        documents = self.process_documents(file_paths)
        context = "\n\n".join(
            [
                (
                    f"[Document: {os.path.basename(path)}]\n{content[:2000]}..."
                    if len(content) > 2000
                    else f"[Document: {os.path.basename(path)}]\n{content}"
                )
                for path, content in documents.items()
            ]
        )

        enhanced_prompt = (
            f"{query}\n\n"
            f"--- CONTENT FROM ATTACHED DOCUMENTS ---\n{context}\n"
            f"--- END OF DOCUMENT CONTENT ---\n\n"
            f"Based on the above documents and the original query, please provide a comprehensive response."
        )
        return enhanced_prompt
