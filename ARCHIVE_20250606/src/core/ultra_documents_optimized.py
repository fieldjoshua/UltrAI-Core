import concurrent.futures
import functools
import hashlib
import io
import logging
import multiprocessing
import os
import pickle
import re
import threading
import time
from collections import OrderedDict
from pathlib import Path
from threading import Lock
from typing import Any, Dict, List, Optional, Set, Tuple, Union

import fitz  # PyMuPDF
import numpy as np
from langchain.text_splitter import RecursiveCharacterTextSplitter

# Only import SentenceTransformer when needed
sentence_transformer_imported = False
faiss_imported = False


# LRU Cache implementation for in-memory caching
class LRUCache:
    """Thread-safe LRU Cache for document chunks and embeddings"""

    def __init__(self, capacity: int = 100):
        """Initialize LRU cache with specified capacity"""
        self.cache = OrderedDict()
        self.capacity = capacity
        self.lock = Lock()

    def get(self, key: str) -> Optional[Any]:
        """Get item from cache, moving it to the end (most recently used)"""
        with self.lock:
            if key not in self.cache:
                return None

            # Move to end (mark as recently used)
            value = self.cache.pop(key)
            self.cache[key] = value
            return value

    def put(self, key: str, value: Any) -> None:
        """Add item to cache, removing oldest item if capacity is exceeded"""
        with self.lock:
            if key in self.cache:
                # Remove existing item first
                self.cache.pop(key)

            # Add new item at the end
            self.cache[key] = value

            # Remove oldest if over capacity
            if len(self.cache) > self.capacity:
                self.cache.popitem(last=False)

    def clear(self) -> None:
        """Clear the cache"""
        with self.lock:
            self.cache.clear()

    def size(self) -> int:
        """Return current cache size"""
        with self.lock:
            return len(self.cache)


class DocumentCache:
    """Enhanced cache for document processing with multi-level caching."""

    def __init__(self, cache_dir: str = ".cache", memory_cache_size: int = 100):
        """Initialize the document cache."""
        self.cache_dir = Path(cache_dir)
        self.cache_dir.mkdir(exist_ok=True)
        self.logger = logging.getLogger(__name__)

        # In-memory LRU cache
        self.memory_cache = LRUCache(capacity=memory_cache_size)

        # Create subdirectories for different cache types
        self.text_cache_dir = self.cache_dir / "text"
        self.chunks_cache_dir = self.cache_dir / "chunks"
        self.embeddings_cache_dir = self.cache_dir / "embeddings"

        for directory in [
            self.text_cache_dir,
            self.chunks_cache_dir,
            self.embeddings_cache_dir,
        ]:
            directory.mkdir(exist_ok=True)

    def _get_cache_path(self, file_path: str, operation: str) -> Path:
        """Generate a unique cache path for a file and operation."""
        file_hash = hashlib.md5(file_path.encode()).hexdigest()

        # Include file modification time for cache invalidation
        if os.path.exists(file_path):
            file_timestamp = os.path.getmtime(file_path)
            cache_key = f"{file_hash}_{int(file_timestamp)}_{operation}"
        else:
            cache_key = f"{file_hash}_{operation}"

        # Select appropriate cache directory based on operation type
        if operation.startswith("text_"):
            cache_dir = self.text_cache_dir
        elif operation.startswith("chunks_"):
            cache_dir = self.chunks_cache_dir
        elif operation.startswith("embeddings_"):
            cache_dir = self.embeddings_cache_dir
        else:
            cache_dir = self.cache_dir

        return cache_dir / f"{cache_key}.pkl"

    def get(self, file_path: str, operation: str) -> Optional[Any]:
        """Get cached result for a file and operation if available."""
        # First check memory cache
        memory_key = f"{file_path}_{operation}"
        memory_result = self.memory_cache.get(memory_key)
        if memory_result is not None:
            self.logger.debug(f"Memory cache hit for {file_path} - {operation}")
            return memory_result

        # If not in memory, check disk cache
        if not os.path.exists(file_path) and not operation.startswith("embeddings_"):
            return None

        cache_path = self._get_cache_path(file_path, operation)
        if cache_path.exists():
            try:
                with open(cache_path, "rb") as f:
                    result = pickle.load(f)

                # Add to memory cache for faster future access
                self.memory_cache.put(memory_key, result)
                self.logger.debug(f"Disk cache hit for {file_path} - {operation}")
                return result
            except Exception as e:
                self.logger.warning(f"Failed to load cache for {file_path}: {e}")
                # Remove corrupted cache file
                try:
                    cache_path.unlink()
                except:
                    pass
                return None
        return None

    def set(self, file_path: str, operation: str, data: Any) -> bool:
        """Cache result for a file and operation."""
        # First store in memory cache
        memory_key = f"{file_path}_{operation}"
        self.memory_cache.put(memory_key, data)

        # Then store on disk
        if not os.path.exists(file_path) and not operation.startswith("embeddings_"):
            return False

        cache_path = self._get_cache_path(file_path, operation)
        try:
            with open(cache_path, "wb") as f:
                pickle.dump(data, f)
            return True
        except Exception as e:
            self.logger.warning(f"Failed to save cache for {file_path}: {e}")
            return False

    def invalidate(self, file_path: str, operation: Optional[str] = None) -> None:
        """Invalidate cache entries for a file."""
        # Remove from memory cache
        if operation:
            memory_key = f"{file_path}_{operation}"
            # Remove specific operation
            self.memory_cache.get(memory_key)  # This will move it to the end
            self.memory_cache.put(memory_key, None)  # Override with None
        else:
            # Clear all operations for this file by creating a new cache
            # This is easier than searching and removing specific keys
            new_cache = LRUCache(self.memory_cache.capacity)
            for k, v in self.memory_cache.cache.items():
                if not k.startswith(file_path):
                    new_cache.put(k, v)
            self.memory_cache = new_cache

        # Remove from disk cache
        file_hash = hashlib.md5(file_path.encode()).hexdigest()
        for cache_dir in [
            self.text_cache_dir,
            self.chunks_cache_dir,
            self.embeddings_cache_dir,
            self.cache_dir,
        ]:
            for cache_file in cache_dir.glob(f"{file_hash}_*"):
                if operation and operation not in cache_file.name:
                    continue
                try:
                    cache_file.unlink()
                except:
                    pass


class UltraDocumentsOptimized:
    def __init__(
        self,
        api_keys: Dict[str, str] = None,
        prompt_templates: Optional[Any] = None,
        rate_limits: Optional[Any] = None,
        output_format: str = "plain",
        enabled_features: Optional[List[str]] = None,
        cache_enabled: bool = True,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        embedding_model: str = "all-MiniLM-L6-v2",
        max_workers: int = None,
        memory_cache_size: int = 100,
    ):
        # Initialize logger
        self.logger = logging.getLogger(__name__)
        self.logger.info("Initializing Optimized Document Processor")

        # Store configuration
        self.api_keys = api_keys or {}
        self.output_format = output_format
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.cache_enabled = cache_enabled
        self.embedding_model_name = embedding_model

        # Set max workers based on CPU cores if not specified
        self.max_workers = max_workers or max(1, multiprocessing.cpu_count() - 1)

        # File processing handlers
        self.supported_formats = {
            ".pdf": self._process_pdf,
            ".txt": self._process_text,
            ".md": self._process_text,
            ".docx": self._process_docx,
            ".doc": self._process_doc,
        }

        # Initialize cache
        if self.cache_enabled:
            self.cache = DocumentCache(memory_cache_size=memory_cache_size)
            self.logger.info(
                f"Document cache initialized with {memory_cache_size} memory slots"
            )

        # Initialize text splitter for chunking with improved settings
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ".", "!", "?", ";", ":", " ", ""],
        )

        # Lazy load embedding model only when needed
        self.embedding_model = None
        self.embedding_dimension = None
        self.index = None
        self.embeddings = []
        self.chunks = []
        self.doc_mapping = {}

        # Thread pool for parallel processing
        self.executor = concurrent.futures.ThreadPoolExecutor(
            max_workers=self.max_workers
        )
        self.logger.info(f"Thread pool initialized with {self.max_workers} workers")

    def _load_embedding_model(self):
        """Lazy-load the embedding model when needed"""
        global sentence_transformer_imported, faiss_imported

        if self.embedding_model is not None:
            return

        try:
            if not sentence_transformer_imported:
                from sentence_transformers import SentenceTransformer

                sentence_transformer_imported = True

            if not faiss_imported:
                import faiss

                faiss_imported = True

            self.embedding_model = SentenceTransformer(self.embedding_model_name)
            self.embedding_dimension = (
                self.embedding_model.get_sentence_embedding_dimension()
            )
            self.logger.info(
                f"Document embedding model loaded: {self.embedding_model_name}"
            )
        except Exception as e:
            self.logger.error(f"Failed to initialize embedding model: {e}")
            self.embedding_model = None

    def _process_pdf(self, file_path: Union[str, Path]) -> str:
        """Extract text from a PDF file using PyMuPDF with optimized processing."""
        try:
            # Check cache first
            if self.cache_enabled:
                cached_text = self.cache.get(str(file_path), "text_pdf")
                if cached_text is not None:
                    self.logger.debug(f"Retrieved PDF text from cache: {file_path}")
                    return cached_text

            # Process PDF if not in cache
            start_time = time.time()
            text_parts = []

            with fitz.open(file_path) as doc:
                # Process in chunks of 10 pages to avoid memory issues with large PDFs
                chunk_size = 10
                total_pages = len(doc)

                for i in range(0, total_pages, chunk_size):
                    end = min(i + chunk_size, total_pages)
                    chunk_text = ""

                    for page_num in range(i, end):
                        page = doc[page_num]
                        # Extract text with metadata
                        page_text = page.get_text()
                        if page_text.strip():
                            chunk_text += f"[Page {page_num + 1}]\n{page_text}\n\n"

                    if chunk_text:
                        text_parts.append(chunk_text)

            text = "".join(text_parts)

            # Cache the result
            if self.cache_enabled:
                self.cache.set(str(file_path), "text_pdf", text)

            processing_time = time.time() - start_time
            self.logger.info(f"Processed PDF in {processing_time:.2f}s: {file_path}")
            return text

        except Exception as e:
            self.logger.error(f"Error processing PDF file {file_path}: {str(e)}")
            raise

    def _process_text(self, file_path: Union[str, Path]) -> str:
        """Extract text from a text file with improved encoding detection."""
        try:
            # Check cache first
            if self.cache_enabled:
                cached_text = self.cache.get(str(file_path), "text_txt")
                if cached_text is not None:
                    self.logger.debug(
                        f"Retrieved text file content from cache: {file_path}"
                    )
                    return cached_text

            # Try UTF-8 first (most common)
            try:
                with open(file_path, "r", encoding="utf-8") as file:
                    text = file.read()
            except UnicodeDecodeError:
                # Try with different encodings if UTF-8 fails
                text = None
                encodings = ["latin-1", "cp1252", "iso-8859-1", "utf-16"]
                for encoding in encodings:
                    try:
                        with open(file_path, "r", encoding=encoding) as file:
                            text = file.read()
                            break
                    except UnicodeDecodeError:
                        continue

                if text is None:
                    self.logger.error(
                        f"Could not decode text file {file_path} with any encoding"
                    )
                    raise ValueError(
                        f"Could not decode text file {file_path} with any supported encoding"
                    )

            # Cache the result
            if self.cache_enabled:
                self.cache.set(str(file_path), "text_txt", text)

            return text

        except Exception as e:
            self.logger.error(f"Error processing text file {file_path}: {str(e)}")
            raise

    def _process_docx(self, file_path: Union[str, Path]) -> str:
        """Extract text from a DOCX file with improved formatting and metadata."""
        try:
            # Check cache first
            if self.cache_enabled:
                cached_text = self.cache.get(str(file_path), "text_docx")
                if cached_text is not None:
                    self.logger.debug(f"Retrieved DOCX text from cache: {file_path}")
                    return cached_text

            from docx import Document

            doc = Document(file_path)

            # Extract paragraphs with better formatting preservation
            sections = []
            current_section = []

            # Process paragraphs
            for para in doc.paragraphs:
                if para.text.strip():  # Skip empty paragraphs
                    # Check if it's a heading (likely section break)
                    if para.style.name.startswith("Heading"):
                        if current_section:
                            sections.append("\n".join(current_section))
                            current_section = []
                        current_section.append(f"## {para.text} ##")
                    else:
                        current_section.append(para.text)

            # Add the last section
            if current_section:
                sections.append("\n".join(current_section))

            # Extract tables
            tables_text = []
            for i, table in enumerate(doc.tables):
                table_rows = []
                table_rows.append(f"[Table {i+1}]")
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_rows.append(" | ".join(row_text))

                if len(table_rows) > 1:  # Only add if table has content
                    tables_text.append("\n".join(table_rows))

            # Combine all text
            all_text = sections + tables_text
            result = "\n\n".join(all_text)

            # Cache the result
            if self.cache_enabled:
                self.cache.set(str(file_path), "text_docx", result)

            return result

        except Exception as e:
            self.logger.error(f"Error processing DOCX file {file_path}: {str(e)}")
            raise

    def _process_doc(self, file_path: Union[str, Path]) -> str:
        """Extract text from a DOC file."""
        try:
            # Check cache first
            if self.cache_enabled:
                cached_text = self.cache.get(str(file_path), "text_doc")
                if cached_text is not None:
                    self.logger.debug(f"Retrieved DOC text from cache: {file_path}")
                    return cached_text

            import textract

            text = textract.process(file_path).decode("utf-8")

            # Cache the result
            if self.cache_enabled:
                self.cache.set(str(file_path), "text_doc", text)

            return text

        except Exception as e:
            self.logger.error(f"Error processing DOC file {file_path}: {str(e)}")
            raise

    def process_document(self, file_path: Union[str, Path]) -> str:
        """Process a document and extract its text content."""
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        file_extension = file_path.suffix.lower()
        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}")

        return self.supported_formats[file_extension](file_path)

    def process_documents(self, file_paths: List[Union[str, Path]]) -> Dict[str, str]:
        """Process multiple documents in parallel and return their text content."""
        results = {}

        # Process documents in parallel using thread pool
        future_to_path = {
            self.executor.submit(self.process_document, file_path): file_path
            for file_path in file_paths
        }

        for future in concurrent.futures.as_completed(future_to_path):
            file_path = future_to_path[future]
            try:
                text = future.result()
                results[str(file_path)] = text
            except Exception as e:
                self.logger.error(f"Error processing {file_path}: {str(e)}")
                results[str(file_path)] = f"Error: {str(e)}"

        return results

    def chunk_text(
        self, text: str, metadata: Dict[str, Any] = None
    ) -> List[Dict[str, Any]]:
        """Split text into semantic chunks with metadata."""
        raw_chunks = self.text_splitter.split_text(text)

        # Add metadata to each chunk
        chunks_with_metadata = []
        for i, chunk in enumerate(raw_chunks):
            chunk_data = {"text": chunk, "chunk_id": i, "total_chunks": len(raw_chunks)}

            # Add any additional metadata
            if metadata:
                chunk_data.update(metadata)

            chunks_with_metadata.append(chunk_data)

        return chunks_with_metadata

    def chunk_document(self, file_path: Union[str, Path]) -> List[Dict[str, Any]]:
        """Process a document, chunk it, and add metadata."""
        try:
            # Extract basic file info
            file_path = Path(file_path)
            file_name = file_path.name
            file_extension = file_path.suffix.lower()
            file_size = file_path.stat().st_size if file_path.exists() else 0

            # Check cache first
            if self.cache_enabled:
                cached_chunks = self.cache.get(
                    str(file_path), f"chunks_{self.chunk_size}_{self.chunk_overlap}"
                )
                if cached_chunks is not None:
                    self.logger.debug(f"Retrieved chunks from cache: {file_path}")
                    return cached_chunks

            # Process document to extract text
            text = self.process_document(file_path)

            # Create metadata
            metadata = {
                "source": str(file_path),
                "filename": file_name,
                "extension": file_extension,
                "file_size": file_size,
                "processed_at": time.time(),
            }

            # Chunk text with metadata
            chunks = self.chunk_text(text, metadata)

            # Cache chunks
            if self.cache_enabled:
                self.cache.set(
                    str(file_path),
                    f"chunks_{self.chunk_size}_{self.chunk_overlap}",
                    chunks,
                )

            return chunks

        except Exception as e:
            self.logger.error(f"Error chunking document {file_path}: {str(e)}")
            raise

    def chunk_documents(
        self, file_paths: List[Union[str, Path]]
    ) -> List[Dict[str, Any]]:
        """Process and chunk multiple documents in parallel."""
        chunked_docs = []

        # Process documents in parallel using thread pool
        future_to_path = {
            self.executor.submit(self.chunk_document, file_path): file_path
            for file_path in file_paths
        }

        for future in concurrent.futures.as_completed(future_to_path):
            file_path = future_to_path[future]
            try:
                chunks = future.result()
                chunked_docs.extend(chunks)
            except Exception as e:
                self.logger.error(f"Error chunking {file_path}: {str(e)}")
                # Continue with other documents

        return chunked_docs

    def embed_text(self, text: str) -> np.ndarray:
        """Generate embedding for a text string."""
        # Lazy load the embedding model if not already loaded
        if self.embedding_model is None:
            self._load_embedding_model()

        if self.embedding_model is None:
            raise ValueError("Embedding model could not be initialized")

        return self.embedding_model.encode(text)

    async def get_relevant_chunks(
        self, query: str, chunks: List[Dict[str, Any]], top_k: int = 5
    ) -> List[Dict[str, Any]]:
        """Get semantically relevant chunks for a query."""
        # Lazy load the embedding model if not already loaded
        if self.embedding_model is None:
            self._load_embedding_model()

        if self.embedding_model is None:
            # Fallback to basic keyword matching if embedding model couldn't be loaded
            self.logger.warning(
                "Embedding model not available, falling back to keyword matching"
            )
            return self._keyword_match_chunks(query, chunks, top_k)

        # Check cache for query embeddings
        query_cache_key = f"query_{hashlib.md5(query.encode()).hexdigest()}"
        if self.cache_enabled:
            cached_embedding = self.cache.get(query_cache_key, "embeddings_query")
            if cached_embedding is not None:
                query_embedding = cached_embedding
            else:
                query_embedding = self.embedding_model.encode(query)
                self.cache.set(query_cache_key, "embeddings_query", query_embedding)
        else:
            query_embedding = self.embedding_model.encode(query)

        # Embed chunks or retrieve from cache
        chunk_embeddings = []
        valid_chunks = []

        import faiss

        for i, chunk in enumerate(chunks):
            chunk_text = chunk["text"]
            chunk_cache_key = f"chunk_{hashlib.md5(chunk_text.encode()).hexdigest()}"

            # Get embedding from cache or compute it
            if self.cache_enabled:
                cached_embedding = self.cache.get(chunk_cache_key, "embeddings_chunk")
                if cached_embedding is not None:
                    chunk_embedding = cached_embedding
                else:
                    chunk_embedding = self.embedding_model.encode(chunk_text)
                    self.cache.set(chunk_cache_key, "embeddings_chunk", chunk_embedding)
            else:
                chunk_embedding = self.embedding_model.encode(chunk_text)

            chunk_embeddings.append(chunk_embedding)
            valid_chunks.append(chunk)

        # Convert to numpy arrays
        chunk_embeddings_array = np.array(chunk_embeddings)

        # Normalize embeddings
        faiss.normalize_L2(query_embedding.reshape(1, -1))
        faiss.normalize_L2(chunk_embeddings_array)

        # Compute similarities and get top_k matches
        similarities = chunk_embeddings_array @ query_embedding.T
        indices = np.argsort(similarities.flatten())[::-1][:top_k]

        # Return top_k relevant chunks with similarity scores
        result = []
        for idx in indices:
            if idx < len(valid_chunks):
                chunk = valid_chunks[idx].copy()
                # Add relevance score (1.0 is best match)
                chunk["relevance"] = float(similarities[idx][0])
                result.append(chunk)

        return result

    def _keyword_match_chunks(
        self, query: str, chunks: List[Dict[str, Any]], top_k: int = 5
    ) -> List[Dict[str, Any]]:
        """Fallback method for finding relevant chunks using keyword matching."""
        # Simple keyword matching algorithm
        query_terms = query.lower().split()
        chunk_scores = []

        for chunk in chunks:
            text = chunk["text"].lower()
            score = 0
            for term in query_terms:
                if term in text:
                    term_count = text.count(term)
                    term_score = term_count / len(text.split())
                    score += term_score

            # Normalize by number of terms
            if query_terms:
                score /= len(query_terms)

            chunk_scores.append((chunk, score))

        # Sort by score (descending)
        chunk_scores.sort(key=lambda x: x[1], reverse=True)

        # Get top_k chunks
        results = []
        for chunk, score in chunk_scores[:top_k]:
            chunk_copy = chunk.copy()
            chunk_copy["relevance"] = score
            results.append(chunk_copy)

        return results

    def cleanup(self):
        """Cleanup resources when done."""
        if hasattr(self, "executor") and self.executor:
            self.executor.shutdown()

        # Clear memory-intensive objects
        self.embedding_model = None
        self.index = None
        self.embeddings = []
        self.chunks = []
        self.doc_mapping = {}

        self.logger.info("Document processor resources cleaned up")
